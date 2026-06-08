from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

MD_PATH = 'visionmate_ieee_paper_draft.md'
OUT_DOC = 'VisionMate_FYP_Complete.docx'


def add_page_number(paragraph):
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)


def style_doc(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)

    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(13)
    h2.font.bold = True


def insert_toc(doc):
    p = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    p._p.append(fld)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    hpara = header.paragraphs[0]
    hpara.text = 'VisionMate — Final Year Project'
    hpara.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    footer = section.footer
    fpara = footer.paragraphs[0]
    fpara.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fpara.text = 'Page '
    add_page_number(fpara)


def read_md():
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        return f.read()


def extract_refs(md):
    # crude extraction: lines starting with [number]
    refs = re.findall(r"\[\d+\].*", md)
    return refs


def add_chapter_placeholder(doc, number, title, lines=None):
    doc.add_heading(f'{number}. {title}', level=1)
    if lines:
        for ln in lines:
            doc.add_paragraph(ln)
    else:
        doc.add_paragraph('[Content to be added]')
    doc.add_page_break()


def create_test_cases_table(doc):
    doc.add_heading('7.1 Extended Test Cases', level=2)
    p = doc.add_paragraph('The following test cases cover core functionality and acceptance criteria. Actual results should be recorded during formal testing.')

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Test Case ID'
    hdr_cells[1].text = 'Module'
    hdr_cells[2].text = 'Test Steps'
    hdr_cells[3].text = 'Test Data'
    hdr_cells[4].text = 'Expected Result'
    hdr_cells[5].text = 'Actual/Status'

    cases = [
        ('TC-01', 'Object Detection', 'Open app and point camera to a person; wait for detection', 'Live camera frame', 'Person detected and distance announced', 'Pending'),
        ('TC-02', 'Face Recognition', 'Enroll a known user; present face to camera', 'Enrolled face template', 'Identity announced as known', 'Pending'),
        ('TC-03', 'Path-State', 'Place an obstacle on walking path then remove', 'Physical obstacle', 'Blocked then cleared announcements', 'Pending'),
        ('TC-04', 'Navigation', 'Start route and follow turn-by-turn prompts', 'Predefined route', 'Turn prompts at correct distances', 'Pending'),
        ('TC-05', 'Emergency SOS', 'Simulate fall and do not cancel within countdown', 'Simulated fall/inertial event', 'SOS published to guardian', 'Pending'),
    ]

    for tc in cases:
        row_cells = table.add_row().cells
        for i in range(6):
            row_cells[i].text = tc[i]

    doc.add_paragraph('\n')


def create_decision_table(doc):
    doc.add_heading('7.2 Decision Table', level=2)
    doc.add_paragraph('Decision logic for SOS escalation and guardian notification. Rows: conditions; columns: rules leading to actions.')

    # Simple decision table as text
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Condition / Rule'
    for i, r in enumerate(['Rule 1', 'Rule 2', 'Rule 3', 'Rule 4']):
        table.rows[0].cells[i+1].text = r

    conds = [
        ('Fall Detected', 'T', 'T', 'F', 'F'),
        ('User Cancels', 'F', 'T', 'T', 'F'),
        ('Battery Critical', 'F', 'F', 'F', 'T'),
    ]
    for idx, c in enumerate(conds, start=1):
        table.rows[idx].cells[0].text = c[0]
        for j in range(4):
            table.rows[idx].cells[j+1].text = c[j+1]

    doc.add_paragraph('Actions: Rule1 -> Prompt and start countdown; Rule2 -> Cancel SOS; Rule3 -> Send SOS immediately; Rule4 -> Publish battery-critical alert.')
    doc.add_page_break()


def create_traceability_matrix(doc):
    doc.add_heading('7.3 Traceability Matrix', level=2)
    doc.add_paragraph('Mapping of high-level requirements (RIDs) to use cases (UCIDs). Use the matrix to ensure coverage between requirements and implemented/tested use cases.')

    # Example smaller matrix: R1..R8 vs UC1..UC8
    reqs = [f'R{n}' for n in range(1,9)]
    ucs = [f'UC{n}' for n in range(1,13)]

    cols = 1 + len(reqs)
    rows = 1 + len(ucs)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    # header
    table.rows[0].cells[0].text = 'UC/Req'
    for i, r in enumerate(reqs):
        table.rows[0].cells[i+1].text = r
    # fill uc rows with placeholders
    for i, uc in enumerate(ucs, start=1):
        table.rows[i].cells[0].text = uc
        for j in range(len(reqs)):
            table.rows[i].cells[j+1].text = ''

    doc.add_paragraph('Note: Fill checkmarks where a requirement maps to a use case. This is an example template; extend as needed to match your full RID/UC set.')
    doc.add_page_break()


def add_results_section(doc):
    doc.add_heading('8. RESULTS/OUTPUT/STATISTICS', level=1)
    doc.add_paragraph('\n')
    # Insert metrics from draft (hard-coded from draft)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Component'
    table.rows[0].cells[1].text = 'Metric'
    table.rows[0].cells[2].text = 'Value'

    rows = [
        ('Object detection', 'Precision', '94.2% (daylight outdoor)'),
        ('Object detection', 'Recall', '88.3% (mixed scenes)'),
        ('Face recognition', 'TAR @ FAR=1%', '96.2%'),
    ]
    for i, r in enumerate(rows, start=1):
        table.rows[i].cells[0].text = r[0]
        table.rows[i].cells[1].text = r[1]
        table.rows[i].cells[2].text = r[2]

    doc.add_paragraph('\n8.1 %completion')
    doc.add_paragraph('Use the matrix & values from 7.3.1 to compute completion: (number of implemented requirements / total requirements) * 100')
    doc.add_paragraph('\n8.2 %accuracy')
    doc.add_paragraph('Use detection precision and validation datasets to compute accuracy per component.')
    doc.add_paragraph('\n8.3 %correctness')
    doc.add_paragraph('Use test outcomes from 7.1 and coverage from 7.3.4 to compute correctness fraction.')
    doc.add_page_break()


def add_conclusion_and_future(doc, md):
    doc.add_heading('9. CONCLUSION', level=1)
    # try to extract conclusion paragraph from md
    m = re.search(r"## XII\. CONCLUSION\n\n(.*?)(?:\n\n##|\Z)", md, re.S)
    if m:
        doc.add_paragraph(m.group(1).strip())
    else:
        doc.add_paragraph('VisionMate demonstrates integrated edge-based assistive intelligence for safer mobility, combining object detection, face recognition, navigation, and guardian workflows.')
    doc.add_page_break()

    doc.add_heading('10. FUTURE WORK', level=1)
    doc.add_paragraph('Suggested future work:')
    doc.add_paragraph('- Larger user studies and longitudinal trials')
    doc.add_paragraph('- Multilingual support and adaptive TTS voices')
    doc.add_paragraph('- Adaptive thresholding and low-light model optimization')
    doc.add_page_break()


def add_bibliography(doc, md):
    doc.add_heading('11. BIBLIOGRAPHY', level=1)
    refs = extract_refs(md)
    if refs:
        for r in refs:
            doc.add_paragraph(r)
    else:
        doc.add_paragraph('Add IEEE-style references here. See visionmate_ieee_paper_draft.md for source references.')
    doc.add_page_break()


def add_appendix(doc):
    doc.add_heading('12. APPENDIX', level=1)
    doc.add_paragraph('12.1 Glossary of Terms')
    terms = [
        ('TTS', 'Text-to-speech: converts text to audible speech.'),
        ('STT', 'Speech-to-text: converts spoken audio to text.'),
        ('YOLO', 'You Only Look Once: real-time object detection family of models.'),
        ('TFLite', 'TensorFlow Lite: optimized runtime for mobile/edge ML models.'),
    ]
    for t in terms:
        doc.add_paragraph(f'- {t[0]}: {t[1]}')

    doc.add_paragraph('\n12.2 Pre-requisites')
    doc.add_paragraph('Development environment and dependencies:')
    # read pubspec
    with open('pubspec.yaml', 'r', encoding='utf-8') as f:
        pub = f.read()
    # crude extract dependencies block
    m = re.search(r'dependencies:\n(.*?)(?:\n\n|\Z)', pub, re.S)
    if m:
        deps_block = m.group(1)
        deps = [ln.strip() for ln in deps_block.splitlines() if ln.strip()]
        for d in deps:
            doc.add_paragraph('- ' + d)
    else:
        doc.add_paragraph('- See pubspec.yaml for Flutter dependencies')

    doc.add_paragraph('\nModel files (assets/models):')
    # list model files
    import os
    models_dir = os.path.join('assets', 'models')
    if os.path.exists(models_dir):
        for f in os.listdir(models_dir):
            doc.add_paragraph('- ' + f)
    else:
        doc.add_paragraph('- models not found in assets/models')

    doc.add_page_break()


def main():
    md = read_md()
    doc = Document()
    style_doc(doc)
    add_header_footer(doc)

    # Title
    p = doc.add_paragraph()
    run = p.add_run('VisionMate: Final Year Project Documentation')
    run.bold = True
    run.font.size = Pt(18)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('\n')
    doc.add_paragraph('Student: ____________________')
    doc.add_paragraph('Supervisor: ____________________')
    doc.add_page_break()

    doc.add_heading('Contents', level=1)
    insert_toc(doc)
    doc.add_page_break()

    # Chapters 1-6: brief extraction
    add_chapter_placeholder(doc, '1', 'INTRODUCTION', extract_lines(md, 'I. INTRODUCTION'))
    add_chapter_placeholder(doc, '2', 'RELATED WORK', extract_lines(md, 'II. RELATED WORK AND GAP ANALYSIS'))
    add_chapter_placeholder(doc, '3', 'SYSTEM ARCHITECTURE', extract_lines(md, 'III. SYSTEM ARCHITECTURE'))
    add_chapter_placeholder(doc, '4', 'METHODS', extract_lines(md, 'IV. METHODS'))
    add_chapter_placeholder(doc, '5', 'GUARDIAN ROLE AND MONITORING WORKFLOW', extract_lines(md, 'V. GUARDIAN ROLE AND MONITORING WORKFLOW'))
    add_chapter_placeholder(doc, '6', 'IMPLEMENTATION DETAILS', extract_lines(md, 'VI. IMPLEMENTATION DETAILS'))

    # Chapter 7: Testing
    doc.add_heading('7. TESTING', level=1)
    create_test_cases_table(doc)
    create_decision_table(doc)
    create_traceability_matrix(doc)

    # Chapter 8: Results
    add_results_section(doc)

    # Chapter 9-10
    add_conclusion_and_future(doc, md)

    # Chapter 11 Bibliography
    add_bibliography(doc, md)

    # Chapter 12 Appendix
    add_appendix(doc)

    doc.save(OUT_DOC)
    print('Saved', OUT_DOC)


# Utility: extract lines between headings

def extract_lines(md, heading):
    # make heading flexible
    try:
        pattern = rf"{re.escape(heading)}(.*?)(?:\n\n[A-Z]{{1,4}}\.|\Z)"
        m = re.search(pattern, md, re.S)
        if m:
            lines = [ln for ln in m.group(1).splitlines() if ln.strip()]
            return lines
    except Exception:
        pass
    return None


if __name__ == '__main__':
    main()
