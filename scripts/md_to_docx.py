from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

md_path = 'visionmate_ieee_paper_draft.md'
docx_path = 'VisionMate_Documentation.docx'


def add_page_number(paragraph):
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)


def add_paragraph_with_bullets(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)


def insert_toc(doc):
    p = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    # Word TOC field: show levels 1-3, with hyperlinks
    fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    p._p.append(fld)


def style_headings(doc):
    # Normal style
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(13)
    h2.font.bold = True

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    hpara = header.paragraphs[0]
    hpara.text = 'VisionMate — Project Documentation'
    hpara.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    footer = section.footer
    fpara = footer.paragraphs[0]
    fpara.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fpara.text = 'Page '
    add_page_number(fpara)


def main():
    doc = Document()
    style_headings(doc)
    add_header_footer(doc)

    # Title page
    t = doc.add_paragraph()
    run = t.add_run('VisionMate: Dual-Zone Smartphone Assistive Intelligence for Safer Mobility of Visually Impaired Users')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('')

    # Abstract placeholder
    doc.add_heading('Abstract', level=1)
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Try to extract the Abstract section from the markdown
    m = re.search(r"## Abstract\n\n(.*?)\n\n## Keywords", content, re.S)
    if m:
        abstract_text = m.group(1).strip()
        doc.add_paragraph(abstract_text)

    # Keywords
    m2 = re.search(r"## Keywords\n\n(.*?)\n\n---", content, re.S)
    if m2:
        doc.add_paragraph('Keywords: ' + m2.group(1).strip())

    doc.add_page_break()

    # Insert TOC placeholder
    doc.add_heading('Contents', level=1)
    insert_toc(doc)
    doc.add_page_break()

    # Now parse rest of markdown and map headings to document headings
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for raw in lines:
        line = raw.rstrip('\n')
        if not line.strip():
            continue

        # map markdown headings to Word headings
        if line.startswith('## '):
            title = line[3:].strip()
            doc.add_heading(title, level=1)
            continue
        if line.startswith('### '):
            title = line[4:].strip()
            doc.add_heading(title, level=2)
            continue
        if line.startswith('#### '):
            title = line[5:].strip()
            doc.add_heading(title, level=3)
            continue

        # horizontal rule -> page break
        if line.startswith('---') or line.startswith('___'):
            doc.add_page_break()
            continue

        # bullet list
        if re.match(r'^\s*[-*]\s+', line):
            text = re.sub(r'^\s*[-*]\s+', '', line)
            add_paragraph_with_bullets(doc, text)
            continue

        # remove markdown emphasis markers
        text = line.replace('**', '')
        text = text.replace('*', '')
        text = text.replace('`', '')

        doc.add_paragraph(text)

    doc.save(docx_path)
    print('Saved', docx_path)


if __name__ == '__main__':
    main()
    main()
