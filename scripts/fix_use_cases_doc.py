from docx import Document
import shutil
import re

SRC = 'Extended use cases.docx'
BACKUP = 'Extended use cases_before_fixes.docx'
OUT = 'Extended use cases.docx'

# Code-derived facts
FALL_FREEFALL = '2.8'
FALL_IMPACT = '24.0'
FALL_PROMPT_SECONDS = '10'
FACE_SIM_THRESHOLD = '0.80'
VOICE_TRIGGER = 'User presses Volume Up twice.'
TRIPLE_PRESS_NAME = 'Favorite Shortcut (Volume Up 3)'

shutil.copyfile(SRC, BACKUP)
print(f'Backup created: {BACKUP}')

doc = Document(SRC)

# Helper to find use case tables: look for tables that contain a cell starting with 'Use Case ID:'
usecase_tables = []
for t in doc.tables:
    found = False
    for row in t.rows:
        for cell in row.cells:
            if cell.text.strip().startswith('Use Case ID'):
                found = True
                break
        if found:
            usecase_tables.append(t)
            break

print(f'Found {len(usecase_tables)} use case tables')

existing_uc_names = []
for t in usecase_tables:
    # extract Use Case ID and Name
    uc_id = ''
    uc_name = ''
    for row in t.rows:
        if len(row.cells) >= 2:
            left = row.cells[0].text.strip()
            right = row.cells[1].text.strip()
            if left.startswith('Use Case ID'):
                uc_id = right
            if left.startswith('Use Case Name'):
                uc_name = right
    existing_uc_names.append((uc_id, uc_name))

# Print found
for uc in existing_uc_names:
    print('Found UC:', uc)

# Fix specific known items
made_changes = False
for t in usecase_tables:
    # normalize cell text lookup
    # find Trigger row and Preconditions row and Normal Flow row
    trigger_row = None
    preconds_row = None
    normalflow_row = None
    description_row = None
    uc_id = ''
    uc_name = ''
    for row in t.rows:
        left = row.cells[0].text.strip()
        right = row.cells[1].text.strip() if len(row.cells)>1 else ''
        if left.startswith('Use Case ID'):
            uc_id = right
        if left.startswith('Use Case Name'):
            uc_name = right
        if left.startswith('Trigger'):
            trigger_row = row
        if left.startswith('Preconditions'):
            preconds_row = row
        if left.startswith('Normal Flow'):
            normalflow_row = row
        if left.startswith('Description'):
            description_row = row
    # Fix UC-17 voice trigger if matches
    if 'UC-17' in uc_id or 'Voice Command' in uc_name:
        # Ensure trigger text matches code
        if trigger_row is not None:
            if VOICE_TRIGGER not in trigger_row.cells[1].text:
                trigger_row.cells[1].text = VOICE_TRIGGER
                made_changes = True
                print('Updated trigger for', uc_id or uc_name)
        # Ensure preconditions include microphone and STT
        if preconds_row is not None:
            pre = preconds_row.cells[1].text
            if 'Microphone' not in pre:
                pre = 'Microphone permission granted.\n' + pre
                preconds_row.cells[1].text = pre
                made_changes = True
                print('Added Microphone precondition for', uc_id or uc_name)
            if 'Speech recognition' not in pre:
                preconds_row.cells[1].text += '\nSpeech recognition available.'
                made_changes = True
                print('Added Speech recognition precondition for', uc_id or uc_name)
    # Fix fall detection table: look for 'Fall' in name or description
    if 'Fall' in uc_name or (description_row and 'fall' in description_row.cells[1].text.lower()):
        # append thresholds note to Description or Normal Flow
        note = f'\n(Note: thresholds in code: freeFall={FALL_FREEFALL}, impact={FALL_IMPACT}, prompt={FALL_PROMPT_SECONDS}s)'
        # append to description cell
        if description_row is not None:
            if FALL_FREEFALL not in description_row.cells[1].text:
                description_row.cells[1].text += note
                made_changes = True
                print('Added fall thresholds note to', uc_id or uc_name)
    # Fix face recognition: ensure similarity threshold note
    if 'Face' in uc_name or (description_row and 'face' in description_row.cells[1].text.lower()):
        if description_row is not None and FACE_SIM_THRESHOLD not in description_row.cells[1].text:
            description_row.cells[1].text += f'\n(Note: similarity threshold used in offline recognition = {FACE_SIM_THRESHOLD})'
            made_changes = True
            print('Added face similarity note to', uc_id or uc_name)

# Add triple-press favorite UC if not present
has_triple = any('favorite' in (name or '').lower() or 'favorite' in (id or '').lower() for id,name in existing_uc_names)
if not has_triple:
    print('Adding Favorite Shortcut UC')
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Use Case ID:'
    table.rows[0].cells[1].text = 'UC-XX'
    def add_row(tbl,label,text):
        row = tbl.add_row()
        row.cells[0].text = label
        row.cells[1].text = text
    add_row(table,'Use Case Name:','Favorite Shortcut (Volume Up 3)')
    add_row(table,'Created By:','')
    add_row(table,'Date Created:','')
    add_row(table,'Actors:','Primary: Visually Impaired User')
    add_row(table,'Description:','Triple-press volume up to activate favorite destination navigation')
    add_row(table,'Trigger:','User presses Volume Up three times.')
    add_row(table,'Preconditions:','Favorite destination configured; GPS permission; network (if needed)')
    add_row(table,'Post conditions:','Navigation to favorite started')
    add_row(table,'Normal Flow:','1. User triple-presses Volume Up.\n2. System confirms favorite and starts navigation.\n3. System gives audio confirmation.')
    add_row(table,'Alternative Flows:','A1: No favorite set -> prompt to set favorite')
    add_row(table,'Exceptions:','Permissions denied, GPS unavailable')
    made_changes = True

if made_changes:
    doc.save(OUT)
    print('Saved changes to', OUT)
else:
    print('No changes needed')
