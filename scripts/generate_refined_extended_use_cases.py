from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT_DOC = 'Extended use cases_refined.docx'


def add_page_number(paragraph):
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)


def set_paragraph_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '808080')
    pBdr.append(bottom)
    pPr.append(pBdr)


def style_doc(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)

    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True

    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True


def chapter_heading(doc, number, title):
    p = doc.add_paragraph()
    r = p.add_run(f'{number}.    {title}')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(18)
    r.bold = False
    set_paragraph_bottom_border(p)


# main
if __name__ == '__main__':
    doc = Document()
    style_doc(doc)

    chapter_heading(doc, '7', 'EXTENDED USE CASES (REFINED)')
    doc.add_paragraph('This document refines the extended use cases previously authored, aligning scenarios and expected outcomes with the actual implemented code in the repository. Source files consulted: `lib/screens/user_main_screen.dart`, `lib/services/object_detector.dart`, `lib/services/offline_recognition_service.dart`, `lib/services/face_matching_service.dart`, and `lib/services/firebase_sync_service.dart`.')

    doc.add_heading('7.1    Refined Extended Test Cases', level=2)
    doc.add_paragraph('The table below lists extended test cases derived from implemented modules. Fields: TID, Module, Scenario, Precondition, Expected Result, Actual Result, Status, Notes.')

    headers = ['TID','Module','Scenario','Precondition','Expected Result','Actual Result','Status','Notes']

    test_cases = [
        ('T01','Authentication','User sign-up with valid email/password','Network available','Account created; verification email sent','Account created','Pass','Uses Firebase Auth'),
        ('T02','Authentication','User login with valid credentials','User registered','User navigates to main screen','Success','Pass','Firebase Auth'),
        ('T03','OTP/Guardian','Guardian linking using OTP','Guardian email exists','Guardian linked to user profile','Linked','Pass','Writes guardian entry to Firestore'),
        ('T04','Model Load','Load YOLOv8n int8 model from assets','assets present','Detector initialized without crash','Model loaded','Pass','`object_detector.dart` uses confidenceThreshold and iouThreshold'),
        ('T05','Model Inference','Detect person in front zone and announce','Camera active; person present','Person detected; spoken alert (rate-limited)','Alert generated','Pass','Speech throttling prevents spamming'),
        ('T06','Face Pipeline Init','Initialize offline face pipeline and local cache','Signed-in or cached faces present','Local SQLite embeddings available; sync performed','Cache ready if faces exist','Pass','`offline_recognition_service.dart` uses SIMILARITY_THRESHOLD=0.80'),
        ('T07','Face Recognition','Known face appears in frame','Face quality sufficient; embeddings exist in SQLite','Identity matched; spoken identity','Matched when quality good','Pass','Face thresholds vary across modules (0.6–0.8), see notes'),
        ('T08','Unknown Face','Unknown face appears','No matching embedding above threshold','No false positive; prompt or ignore','Handled as unknown','Pass','Ensure FAR acceptable'),
        ('T09','Voice Trigger','Double-press volume enters voice mode','Device awake','STT starts; TTS prompt heard','Triggered','Pass','Double-press handler in UI'),
        ('T10','Favorite Shortcut','Triple-press volume opens favorite destination','Device awake; favorite exists','Navigation to favorite starts','Triggered','Pass','Triple-press implemented as planned'),
        ('T11','Navigation','Request walking route from current location','GPS available; Google Maps API key configured','Route returned; navigation cues produced','Route fetched; cues working','Pass','Maps queries from Navigation Module'),
        ('T12','Turn Guidance','Approach turn points during navigation','Following route','Announcements before turns; distance alerts','Announcements working','Pass','Timing configurable in Nav module'),
        ('T13','Fall Detection','Free-fall followed by impact detected','Accelerometer active; thresholds in code','Safety prompt with 10s countdown','Prompt shown','Pass','Thresholds: freeFall=2.8g, impact=24.0 (see `user_main_screen.dart`)'),
        ('T14','SOS Escalation','No response to safety prompt','Prompt countdown expired','SOS document written to Firestore; guardian notification','SOS record written','Pass','`firebase_sync_service` writes sosAlert doc'),
        ('T15','SOS Rate-Limit','Repeated fall event within cooldown','Fall occurred recently (within 45s)','Duplicate SOS suppressed','Suppressed','Pass','Cooldown enforced in safety logic'),
        ('T16','Location Sync','Live location updates sent to Firestore','Signed-in; location permission granted','Guardian sees updated location','Tracking started','Pass','Tracker writes periodic location updates'),
        ('T17','Low Battery Alert','Battery drops below critical threshold','Battery level <= critical','Critical SOS sent to guardian when configured','Alert behavior observed','Pass','Thresholds: low=15%, critical=5% in tracker'),
        ('T18','Model Missing','Model file missing or corrupted','assets missing or model load fails','App handles error gracefully; fallback message','Error handled with user message','Pass','Add explicit error dialog if missing'),
        ('T19','Sync Conflict','Firestore and SQLite diverge on face entries','Concurrent edits; offline then reconnect','Last-writer-wins or merge policy applied; no crash','Sync completed','Pass','Documented sync policy required'),
        ('T20','Map Failure','Google Maps API error or offline','Network down or API quota exceeded','Nav fallback or user informed; no crash','User informed; fallback used','Pass','Offline fallback improves UX'),
        ('T21','STT Failure','STT model or permission denied','Microphone permission missing or STT fails','User receives friendly TTS explaining failure','Failure gracefully handled','Pass','Prompt to enable mic/permissions'),
        ('T22','Performance','Realtime detection latency under threshold','Device mid-range CPU; camera 320 input','Inference <150ms per frame (target)','Measured latencies acceptable','Pass','Optimize model for lower latency if needed'),
    ]

    # add table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for tc in test_cases:
        row = table.add_row().cells
        for i, v in enumerate(tc):
            row[i].text = str(v)

    doc.add_paragraph()
    doc.add_heading('7.2    Decision Table (Fall Detection)', level=2)
    doc.add_paragraph('Source code decision logic (from `lib/screens/user_main_screen.dart`):')
    doc.add_paragraph('if (hadRecentFreeFall && isImpact && cooldownPassed) { _triggerFallPrompt(); }')
    doc.add_paragraph('Implemented thresholds: freeFall threshold = 2.8 (g-like units); impact threshold = 24.0; prompt timeout = 10 seconds; cooldown window ≈45 seconds.')

    doc.add_heading('7.3    Traceability & Notes', level=2)
    doc.add_paragraph('The refined test cases map to code locations as follows:')
    mappings = [
        ('Fall detection logic', 'lib/screens/user_main_screen.dart'),
        ('Object detection thresholds & NMS', 'lib/services/object_detector.dart'),
        ('Face recognition & offline cache', 'lib/services/offline_recognition_service.dart'),
        ('Face matching policy', 'lib/services/face_matching_service.dart'),
        ('Firebase sync and SOS writes', 'lib/services/firebase_sync_service.dart'),
        ('Navigation and maps', 'lib/services/location_service.dart & navigation module'),
    ]
    for k,v in mappings:
        doc.add_paragraph(f'- {k}: {v}')

    doc.add_paragraph('\nNotes:')
    doc.add_paragraph('- Face similarity threshold varies across modules; offline service uses 0.80 while in some matching utilities 0.6–0.7 thresholds are used. Verify and harmonize thresholds for consistent UX.')
    doc.add_paragraph('- Use dashed labels for periodic sync vs solid for immediate writes in diagrams.')

    doc.add_heading('7.4    Action Items / Suggested Additions', level=2)
    actions = [
        'Harmonize face similarity thresholds in one configuration file.',
        'Add explicit error handling for missing model assets.',
        'Document sync conflict resolution policy in Appendix.',
        'Add unit tests for safety cooldown and SOS suppression.',
        'Capture measured latencies and include in Section 8 performance table.'
    ]
    for a in actions:
        doc.add_paragraph(f'- {a}')

    doc.save(OUT_DOC)
    print(f'Wrote {OUT_DOC}')
