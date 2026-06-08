from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

md_path = 'visionmate_ieee_paper_draft.md'
docx_path = 'VisionMate_FYP_Document.docx'


def add_page_number(paragraph):
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)


def insert_toc(doc):
    p = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    p._p.append(fld)


def style_doc(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)

    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(13)
    h2.font.bold = True


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    hpara = header.paragraphs[0]
    hpara.text = 'VisionMate — Final Year Project'
    hpara.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    footer = section.footer
    fpara = footer.paragraphs[0]
    fpara.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fpara.text = 'Page '
    add_page_number(fpara)


def read_markdown():
    with open(md_path, 'r', encoding='utf-8') as f:
        return f.read()


def add_chapter(doc, number, title, content_lines=None):
    doc.add_heading(f'{number}. {title}', level=1)
    if content_lines:
        for line in content_lines:
            if not line.strip():
                continue
            doc.add_paragraph(line)
    else:
        doc.add_paragraph('''[PLACEHOLDER] Add content for this chapter.\n
Explain methodology, include figures and tables as needed.\n
Use the following placeholders for images: [FIGURE: figure_name.png] and for tables: [TABLE: table_name].''')
    doc.add_page_break()


def extract_section_text(md, section_title):
    # crude extraction: find heading and grab until next '## ' heading
    import re
    pattern = rf"##+\s*{re.escape(section_title)}(.*?)(?:\n##+\s|\Z)"
    m = re.search(pattern, md, re.S | re.I)
    if m:
        text = m.group(1).strip()
        lines = [ln for ln in text.splitlines() if ln.strip()]
        return lines
    return None


def main():
    doc = Document()
    style_doc(doc)
    add_header_footer(doc)

    # Title page
    p = doc.add_paragraph()
    run = p.add_run('VisionMate: Dual-Zone Smartphone Assistive Intelligence for Safer Mobility of Visually Impaired Users')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('\n')
    p2 = doc.add_paragraph()
    p2.add_run('Student: ____________________\nSupervisor: ____________________\nDepartment: ____________________\nUniversity: ____________________\nDate: ____________________')
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # Contents (TOC)
    doc.add_heading('Contents', level=1)
    insert_toc(doc)
    doc.add_page_break()

    md = read_markdown()

    # Chapters as per the provided FYP template picture
    # 7. Testing
    testing_lines = extract_section_text(md, 'VII. EXPERIMENTAL DESIGN') or extract_section_text(md, 'VII. EXPERIMENTAL DESIGN')
    # We'll include subsections placeholders
    add_chapter(doc, '1', 'INTRODUCTION', extract_section_text(md, 'I. INTRODUCTION'))
    add_chapter(doc, '2', 'RELATED WORK AND GAP ANALYSIS', extract_section_text(md, 'II. RELATED WORK AND GAP ANALYSIS'))
    add_chapter(doc, '3', 'SYSTEM ARCHITECTURE', extract_section_text(md, 'III. SYSTEM ARCHITECTURE'))
    add_chapter(doc, '4', 'METHODS/IMPLEMENTATION DETAILS', extract_section_text(md, 'IV. METHODS') or extract_section_text(md, 'VI. IMPLEMENTATION DETAILS'))
    add_chapter(doc, '5', 'IMPLEMENTATION AND GUARDIAN WORKFLOW', extract_section_text(md, 'V. GUARDIAN ROLE AND MONITORING WORKFLOW'))
    add_chapter(doc, '6', 'EXPERIMENTAL DESIGN', extract_section_text(md, 'VII. EXPERIMENTAL DESIGN') or extract_section_text(md, 'VII. EXPERIMENTAL DESIGN'))

    # Testing chapter (7)
    doc.add_heading('7. TESTING', level=1)
    doc.add_paragraph('[7.1] Extended Test Cases: Include test case tables, steps, expected vs actual, pass/fail, notes.')
    doc.add_paragraph('[7.2] Decision Table: Insert decision tables and code snippets as images or formatted text.')
    doc.add_paragraph('[7.3] Traceability Matrix: Insert a traceability matrix image or a table linking requirements to test cases and prototypes.')
    doc.add_page_break()

    # Results/Output/Statistics (8)
    add_chapter(doc, '8', 'RESULTS/OUTPUT/STATISTICS', extract_section_text(md, 'VIII. RESULTS AND ANALYSIS') or extract_section_text(md, 'VIII. RESULTS AND ANALYSIS'))

    # Conclusion (9)
    add_chapter(doc, '9', 'CONCLUSION', extract_section_text(md, '## XII. CONCLUSION') or extract_section_text(md, 'XII. CONCLUSION') )

    # Future Work (10)
    add_chapter(doc, '10', 'FUTURE WORK', extract_section_text(md, 'Future work') or ['Suggestions: multilingual support, larger user studies, adaptive thresholding.'])

    # Bibliography (11)
    doc.add_heading('11. BIBLIOGRAPHY', level=1)
    # Try to append references block
    refs = extract_section_text(md, '## REFERENCES') or extract_section_text(md, 'REFERENCES')
    if refs:
        for line in refs:
            doc.add_paragraph(line)
    else:
        doc.add_paragraph('Include IEEE-style citations here.')
    doc.add_page_break()

    # Appendix (12)
    doc.add_heading('12. APPENDIX', level=1)
    doc.add_paragraph('12.1 Glossary of terms: [Add important terms and definitions].')
    doc.add_paragraph('12.2 Pre-requisites: Development and deployment setup, external system dependencies, environment setup instructions.')
    doc.add_page_break()

    # Save
    doc.save(docx_path)
    print('Saved', docx_path)


if __name__ == '__main__':
    main()
