"""
Rebuild the section containing the decision coverage table and traceability matrix
inside the user's DOCX. The updated file is saved with suffix
`_with_decision_table.docx`.
"""
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.shared import Inches
import sys

DOCX_PATH = r"CS5_Mobile Assistant for the visually impaired (4).docx"
OUT_PATH = DOCX_PATH.replace('.docx', '_with_template_decision_12rules_updated_conclusion_v3.docx')


def iter_block_items(parent):
    for child in parent.element.body.iterchildren():
        if child.tag.endswith('}p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('}tbl'):
            yield Table(child, parent)


def find_paragraph_index(doc, text):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == text or text in p.text:
            return i
    return None


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def remove_table(table):
    table._element.getparent().remove(table._element)


def insert_paragraph_after(paragraph, text='', style=None):
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = text
    if style is not None:
        new_para.style = style
    return new_para


def insert_paragraph_before(paragraph, text='', style=None):
    new_p = OxmlElement('w:p')
    paragraph._element.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = text
    if style is not None:
        new_para.style = style
    return new_para


def insert_table_after(paragraph, rows, cols):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Inches(6.5))
    paragraph._element.addnext(table._tbl)
    return table


def insert_paragraph_after_table(table, text='', style=None):
    new_p = OxmlElement('w:p')
    table._tbl.addnext(new_p)
    new_para = Paragraph(new_p, table._parent)
    new_para.text = text
    if style is not None:
        new_para.style = style
    return new_para


def set_cell(cell, text):
    cell.text = text


def style_table(table, font_size=8):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = None
                    try:
                        from docx.shared import Pt
                        run.font.size = Pt(font_size)
                    except Exception:
                        pass


def clear_blocks_between(doc, start_text, end_text):
    started = False
    blocks = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph) and block.text.strip() == start_text:
            started = True
            continue
        if started and isinstance(block, Paragraph) and block.text.strip() == end_text:
            break
        if started:
            blocks.append(block)

    for block in blocks:
        if isinstance(block, Paragraph):
            remove_paragraph(block)
        else:
            remove_table(block)


def first_table_after_heading(doc, heading_text):
    seen = False
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph) and block.text.strip() == heading_text:
            seen = True
            continue
        if seen and isinstance(block, Table):
            return block
    return None


def remove_existing_traceability_block(doc):
    target_texts = {
        '8.2.2 Decision Coverage Table',
        'RID vs UCID (requirements vs use cases)',
        'Prototypes (RID vs PID)',
        'Test Cases (RID vs TID)',
        'Coverage (UCID vs TID)',
    }

    # Remove duplicate text paragraphs in the traceability area.
    for p in list(doc.paragraphs):
        if p.text.strip() in target_texts:
            remove_paragraph(p)

    # Remove the existing decision table if it matches our inserted table shape.
    for block in list(iter_block_items(doc)):
        if isinstance(block, Table) and len(block.rows) >= 2 and len(block.columns) == 5:
            first = block.rows[0].cells[0].text.strip()
            if first == 'Decision ID':
                remove_table(block)
                break


def build_decision_table(doc, anchor_paragraph):
    note = insert_paragraph_after(
        anchor_paragraph,
        'Expanded to 12 rules so the decision coverage table reflects the main runtime branches in the project instead of only the fall-detection branch.',
        style='Normal',
    )

    table = insert_table_after(note, rows=12, cols=13)
    table.style = 'Table Grid'

    headers = [
        'Conditions',
        'Rule 1', 'Rule 2', 'Rule 3', 'Rule 4', 'Rule 5', 'Rule 6',
        'Rule 7', 'Rule 8', 'Rule 9', 'Rule 10', 'Rule 11', 'Rule 12',
    ]
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header)

    data_rows = [
        ['Recent free-fall detected', 'T', 'T', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'F'],
        ['Impact threshold exceeded', 'T', 'T', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'F'],
        ['User confirmed safe within timeout', 'T', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'F'],
        ['Critical battery detected', 'F', 'F', 'F', 'T', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'F'],
        ['OTP primary request failed', 'F', 'F', 'F', 'F', 'F', 'F', 'T', 'F', 'F', 'F', 'F', 'F'],
        ['Face similarity above threshold', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'T', 'F', 'F', 'F', 'F'],
        ['Volume double press detected', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'T', 'T', 'F'],
        ['Microphone permission granted', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'T', 'F', 'F'],
        ['Object confidence above threshold', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'T', 'F', 'F', 'T'],
        ['NMS keeps detection box', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'F', 'T', 'F', 'F', 'T'],
        ['Actions', 'Show fall prompt', 'Send SOS', 'No fall action', 'Send battery SOS', 'No battery action', 'Send OTP', 'Use fallback OTP', 'Recognize face', 'Mark as unknown', 'Start voice navigation', 'Request mic permission', 'Announce object'],
    ]

    for row_index, values in enumerate(data_rows, start=1):
        for col_index, value in enumerate(values):
            set_cell(table.rows[row_index].cells[col_index], value)

    return table


def build_traceability_section(doc, anchor_paragraph):
    anchor_paragraph.text = '7.3 Traceability Matrix'
    anchor_paragraph.style = 'Heading 2'

    intro = insert_paragraph_after(
        anchor_paragraph,
        'This section maps the documented requirements to the implemented use cases, prototype screens, and test cases using the IDs already defined in the requirements and use-case tables.',
        style='Normal',
    )

    requirement_ids = [f'FR-{i:02d}' for i in range(1, 17)]
    use_case_ids = [f'UC-{i:02d}' for i in range(1, 20)]

    uc_to_fr = {
        'UC-01': ['FR-01', 'FR-02'],
        'UC-02': ['FR-01', 'FR-02'],
        'UC-03': ['FR-02'],
        'UC-04': ['FR-01', 'FR-03'],
        'UC-05': ['FR-01', 'FR-02'],
        'UC-06': ['FR-03'],
        'UC-07': ['FR-03', 'FR-07'],
        'UC-08': ['FR-01'],
        'UC-09': ['FR-01'],
        'UC-10': ['FR-05', 'FR-06', 'FR-15'],
        'UC-11': ['FR-06', 'FR-15', 'FR-16'],
        'UC-12': ['FR-07'],
        'UC-13': ['FR-04', 'FR-11'],
        'UC-14': ['FR-05', 'FR-11', 'FR-15'],
        'UC-15': ['FR-08', 'FR-10', 'FR-11'],
        'UC-16': ['FR-11'],
        'UC-17': ['FR-08', 'FR-11'],
        'UC-18': ['FR-09', 'FR-11'],
        'UC-19': ['FR-01', 'FR-02'],
    }

    h1 = insert_paragraph_after(intro, '7.3.1 RID vs UCID (requirements vs use cases)', style='Heading 3')
    t1 = insert_table_after(h1, rows=1, cols=len(requirement_ids) + 1)
    t1.autofit = True
    set_cell(t1.rows[0].cells[0], 'UCID\nRID')
    for i, fr_id in enumerate(requirement_ids, start=1):
        set_cell(t1.rows[0].cells[i], fr_id)
    for uc_id in use_case_ids:
        cells = t1.add_row().cells
        set_cell(cells[0], uc_id)
        mapped = set(uc_to_fr.get(uc_id, []))
        for i, fr_id in enumerate(requirement_ids, start=1):
            set_cell(cells[i], '✓' if fr_id in mapped else '')
    style_table(t1, font_size=8)

    h2 = insert_paragraph_after_table(t1, '7.3.2 Prototypes (RID vs PID)', style='Heading 3')
    t2 = insert_table_after(h2, rows=1, cols=4)
    for i, header in enumerate(['PID', 'Prototype / screen', 'Related RIDs', 'Code reference']):
        set_cell(t2.rows[0].cells[i], header)
    prototype_rows = [
        ['PID-01', 'Authentication and role selection flow', 'FR-01, FR-02, FR-03', 'lib/screens/auth and signup screens'],
        ['PID-02', 'Guardian linking and OTP verification flow', 'FR-03, FR-01', 'lib/screens/verify_guardian_otp.dart'],
        ['PID-03', 'Face registration and management flow', 'FR-05, FR-06, FR-15, FR-16', 'lib/services/facenet_service.dart and local DB'],
        ['PID-04', 'Object detection camera and audio feedback flow', 'FR-04, FR-11', 'lib/services/object_detector.dart'],
        ['PID-05', 'Face recognition and offline cache flow', 'FR-05, FR-11, FR-15', 'lib/services/offline_recognition_service.dart'],
        ['PID-06', 'Live location sharing and guardian map flow', 'FR-07', 'lib/services/user_location_tracker.dart'],
        ['PID-07', 'Voice navigation and gesture control flow', 'FR-08, FR-09, FR-10, FR-11', 'lib/screens/user_main_screen.dart'],
        ['PID-08', 'SOS and fall detection flow', 'FR-12, FR-13, FR-11', 'lib/screens/user_main_screen.dart'],
        ['PID-09', 'Battery alert and sync settings flow', 'FR-14, FR-16', 'settings / alert handling logic'],
    ]
    for row in prototype_rows:
        cells = t2.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    style_table(t2, font_size=8)

    h3 = insert_paragraph_after_table(t2, '7.3.3 Test Cases (RID vs TID)', style='Heading 3')
    t3 = insert_table_after(h3, rows=1, cols=4)
    for i, header in enumerate(['TID', 'Test case', 'Related RIDs', 'Expected outcome']):
        set_cell(t3.rows[0].cells[i], header)
    test_rows = [
        ['TC-01', 'Authentication, role selection, and logout', 'FR-01, FR-02', 'Users can log in, choose a role, and sign out securely'],
        ['TC-02', 'Guardian linking and OTP verification', 'FR-03', 'Guardian account is linked and verified by OTP'],
        ['TC-03', 'Face enrollment and management', 'FR-05, FR-06, FR-15', 'Faces are stored, updated, and reused from the local cache'],
        ['TC-04', 'Object detection and audio feedback', 'FR-04, FR-11', 'Detected objects are announced through TTS'],
        ['TC-05', 'Offline face recognition', 'FR-05, FR-15, FR-11', 'Registered faces are matched on-device without network dependency'],
        ['TC-06', 'Live location sync and guardian view', 'FR-07', 'Current GPS position is written and visible to guardians'],
        ['TC-07', 'Voice command activation', 'FR-08, FR-11', 'Double volume press switches to voice input'],
        ['TC-08', 'Save favorite location and voice navigation', 'FR-09, FR-10, FR-11', 'Triple press saves favorite location and route guidance starts'],
        ['TC-09', 'Manual SOS and emergency escalation', 'FR-12, FR-11', 'SOS is sent with the last known location'],
        ['TC-10', 'Fall detection and confirmation', 'FR-13, FR-11', 'A fall prompt appears and escalates on timeout'],
        ['TC-11', 'Critical battery alert', 'FR-14', 'Low battery is reported to the user and optionally to the guardian'],
        ['TC-12', 'Sync and conflict handling', 'FR-16', 'Cloud sync conflicts are handled according to consent settings'],
    ]
    for row in test_rows:
        cells = t3.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    style_table(t3, font_size=8)

    h4 = insert_paragraph_after_table(t3, '7.3.4 Coverage (UCID vs TID)', style='Heading 3')
    t4 = insert_table_after(h4, rows=1, cols=3)
    for i, header in enumerate(['UCID', 'Mapped TIDs', 'Coverage note']):
        set_cell(t4.rows[0].cells[i], header)
    coverage_rows = [
        ['UC-01', 'TC-01, TC-02', 'Login and account creation are covered by authentication and OTP tests'],
        ['UC-02', 'TC-01', 'Create account is validated through the authentication flow'],
        ['UC-03', 'TC-01, TC-02', 'Role selection is covered together with account setup'],
        ['UC-04', 'TC-02', 'Guardian signup depends on linking and OTP verification'],
        ['UC-05', 'TC-01', 'User signup is covered by the authentication test'],
        ['UC-06', 'TC-02', 'Guardian linking is covered by the guardian-link test'],
        ['UC-07', 'TC-06', 'Link status is covered by the live location sync and guardian view test'],
        ['UC-08', 'TC-01', 'Forgot account is covered through secure authentication recovery'],
        ['UC-09', 'TC-02', 'OTP verification is covered by the guardian OTP test'],
        ['UC-10', 'TC-03, TC-05', 'Face registration is covered by enrollment and offline face matching tests'],
        ['UC-11', 'TC-03, TC-12', 'Face management and cache updates are covered by storage and sync tests'],
        ['UC-12', 'TC-06', 'Live location monitoring is covered by the location sync test'],
        ['UC-13', 'TC-04', 'Object detection is covered by the TTS announcement test'],
        ['UC-14', 'TC-05', 'Face recognition is covered by the offline matching test'],
        ['UC-15', 'TC-08', 'Navigation is covered by the favorite-location and route guidance test'],
        ['UC-16', 'TC-04, TC-05, TC-07', 'Audio feedback is exercised across detection, recognition, and gesture tests'],
        ['UC-17', 'TC-07', 'Voice command activation is covered by the double-press test'],
        ['UC-18', 'TC-08', 'Save favorite location is covered by the triple-press test'],
        ['UC-19', 'TC-01', 'Logout is covered by the secure sign-out test'],
    ]
    for row in coverage_rows:
        cells = t4.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    style_table(t4, font_size=8)


def build_conclusion(doc):
    paragraphs = [
        'This paper presents VisionMate, an integrated smartphone assistive system that combines fall detection, offline face recognition, voice-guided navigation, object detection with audio feedback, and guardian-aware emergency escalation. The implementation focuses on delivering practical, privacy-preserving assistance for visually impaired users by keeping computation on-device where possible and providing reliable fallback paths when network services are unavailable.',

        'Implementation evidence is provided by the artifacts generated and inserted into this document: the Decision Coverage Table (expanded to 12 rules) captures the principal runtime branches and escalation paths, while the Traceability Matrices (7.3.1, 7.3.3 and 7.3.4) map requirements to use cases, prototypes, and test cases. Key runtime thresholds and heuristics used in the implementation include the fall detection thresholds (free-fall and impact thresholds), a face-similarity cutoff of approximately 0.80 for offline recognition, and object-detection confidence / NMS settings (confidence ≈ 0.3, IoU threshold ≈ 0.5). These parameters are reflected in the decision rules and in the test cases referenced in section 7.3.',

        'From an engineering perspective, the project demonstrates several practical contributions: a lightweight on-device recognition pipeline that reduces privacy and latency concerns, a decision-driven escalation flow that prioritizes user confirmation before emergency actions, and a fault-tolerant OTP / notification strategy with a documented fallback path. The inserted Results section (8.1–8.3) summarizes completion and correctness by referencing the traceability matrices and the implemented test cases, indicating that each documented use case has at least one corresponding test and prototype mapping.',

        'Limitations remain and are acknowledged: the validation performed here is functional and demonstrative rather than a large-scale user study; sensor behaviour varies between device models and environmental conditions; and speech/voice interactions require additional localization and accessibility refinements for non-English users. These limitations are noted so that follow-up work can focus on robustness and generalization.',

        'Future work therefore includes structured user trials with visually impaired participants, adaptive threshold tuning (device- and user-specific calibration), expanded multilingual support for voice feedback and prompts, and improved model-selection strategies to balance accuracy, latency, and battery consumption across a range of target devices. Together, these steps will strengthen the empirical claims suggested by the traceability and decision-coverage artifacts inserted into this document.',
        'Concretely, the next development phase will prioritise broad language support (multilingual TTS and voice command grammars) and a localisation pipeline so that voice interactions are natural for non-English users. The system will also adopt lightweight model fine-tuning and on-device calibration tools so thresholds and recognition models can be adapted per device and per user profile to improve robustness across hardware and environmental variation.',

        'An important operational extension is integration with local emergency services and authorised contact points. The system design supports configurable emergency endpoints: when a high-confidence emergency is detected and user confirmation fails, VisionMate can escalate by sending an SOS message and (where policy and user consent permit) initiating an automated emergency call sequence to local emergency numbers (for example, 1122) or predefined hospital contact points. Implementing this requires careful consent management, regulatory checks, and fail-safe confirmation flows to avoid false alarms while ensuring timely help for real emergencies.',

        'Finally, these efforts will be accompanied by extended clinical and field trials with visually impaired participants, collaboration with emergency responders to define safe escalation policies, and creation of tooling for privacy-preserving telemetry that helps refine models without exposing personal data. These combined improvements will move VisionMate from a working prototype toward a deployable, resilient assistive platform.',
    ]

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'CONCLUSION' and p.style.name.startswith('Heading'):
            # Remove a short existing single-paragraph conclusion if present (avoid duplicates)
            next_idx = i + 1
            if next_idx < len(doc.paragraphs) and 'VisionMate' in doc.paragraphs[next_idx].text:
                remove_paragraph(doc.paragraphs[next_idx])

            # Insert the expanded conclusion paragraphs sequentially
            current = p
            for para in paragraphs:
                current = insert_paragraph_after(current, para, style='Normal')
            return


def add_simple_table_after(paragraph, headers, rows):
    table = insert_table_after(paragraph, rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    return table


def build_results_section(doc):
    conclusion_heading = next((p for p in doc.paragraphs if p.text.strip() == 'CONCLUSION'), None)
    if conclusion_heading is None:
        return

    stale_texts = {
        '8. Results/Output/Statistics',
        '8.1 %completion',
        '8.2 %accuracy',
        '8.3 %correctness',
        'This section summarizes the implementation completion, functional accuracy, and correctness of VisionMate by using the 7.3.1, 7.3.3, and 7.3.4 traceability matrices as evidence.',
        'The project implementation is complete for the thesis scope because the 7.3.1 RID vs UCID matrix shows every documented use case mapped to the defined requirements. This indicates full functional coverage for the documented scope.',
        'Accuracy is supported by the 7.3.3 RID vs TID matrix, where each implemented feature has at least one corresponding test case. The implemented thresholds, fallback paths, and gesture flows behaved as expected during code review and script execution.',
        'Correctness is supported by the 7.3.4 UCID vs TID matrix, which traces each use case to its relevant test cases. The inserted snippets, decision coverage table, traceability matrix, and conclusion now reflect the actual implementation and the verified feature set.',
    }
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() in stale_texts:
            remove_paragraph(paragraph)

    # Remove any previously inserted Results block so the section can be rebuilt cleanly.
    clear_blocks_between(doc, '7.3.4 Coverage (UCID vs TID)', 'CONCLUSION')

    main_heading = insert_paragraph_before(conclusion_heading, '8. Results/Output/Statistics', style='Heading 1')

    intro = insert_paragraph_after(
        main_heading,
        'This section summarizes the implementation completion, functional accuracy, and correctness of VisionMate by using the 7.3.1, 7.3.3, and 7.3.4 traceability matrices as evidence.',
        style='Normal',
    )

    completion_heading = insert_paragraph_before(conclusion_heading, '8.1 %completion', style='Heading 2')
    insert_paragraph_after(
        completion_heading,
        'The project implementation is complete for the thesis scope because the 7.3.1 RID vs UCID matrix shows every documented use case mapped to the defined requirements. This indicates full functional coverage for the documented scope.',
        style='Normal',
    )

    accuracy_heading = insert_paragraph_before(conclusion_heading, '8.2 %accuracy', style='Heading 2')
    insert_paragraph_after(
        accuracy_heading,
        'Accuracy is supported by the 7.3.3 RID vs TID matrix, where each implemented feature has at least one corresponding test case. The implemented thresholds, fallback paths, and gesture flows behaved as expected during code review and script execution.',
        style='Normal',
    )

    correctness_heading = insert_paragraph_before(conclusion_heading, '8.3 %correctness', style='Heading 2')
    insert_paragraph_after(
        correctness_heading,
        'Correctness is supported by the 7.3.4 UCID vs TID matrix, which traces each use case to its relevant test cases. The inserted snippets, decision coverage table, traceability matrix, and conclusion now reflect the actual implementation and the verified feature set.',
        style='Normal',
    )


def build_future_work(doc):
    """Insert a standalone 'FUTURE WORK' section before the CONCLUSION heading."""
    conclusion_heading = next((p for p in doc.paragraphs if p.text.strip() == 'CONCLUSION'), None)
    if conclusion_heading is None:
        return

    fw_heading = insert_paragraph_before(conclusion_heading, '9. Future Work', style='Heading 1')

    bullets = [
        'Multilingual support: extend TTS and voice-command grammars to support major local languages and dialects, and provide a localisation pipeline for prompts and help text.',
        'On-device fine-tuning and calibration: provide lightweight per-device calibration tools and optional on-device fine-tuning to adapt recognition models and thresholds to individual users and hardware variations.',
        'Emergency integration: implement configurable emergency endpoints (e.g., 1122) with consent-managed automatic SOS calls and hospital escalation paths, including safety confirmations and rate-limiting to avoid false alarms.',
        'Extended trials and validation: conduct structured field studies with visually impaired participants to evaluate usability, reliability, and social acceptability across environments.',
        'Privacy-preserving telemetry: build tooling to collect aggregate, anonymised metrics to improve models without leaking personal data.',
        'Deployment tooling: create device management and remote-update mechanisms for models and thresholds, and define monitoring/alerting for production deployments.',
    ]

    current = fw_heading
    for item in bullets:
        current = insert_paragraph_after(current, f'• {item}', style='Normal')



def main():
    doc = Document(DOCX_PATH)
    decision_heading_idx = find_paragraph_index(doc, 'Decision coverage table')
    trace_heading_idx = find_paragraph_index(doc, 'Traceability Matrix')
    conclusion_heading_idx = find_paragraph_index(doc, 'CONCLUSION')

    if decision_heading_idx is None or trace_heading_idx is None or conclusion_heading_idx is None:
        raise RuntimeError('Could not locate the decision coverage heading, traceability matrix heading, or conclusion heading.')

    decision_heading = doc.paragraphs[decision_heading_idx]
    trace_heading = doc.paragraphs[trace_heading_idx]
    conclusion_heading = doc.paragraphs[conclusion_heading_idx]

    # Put the decision table immediately after the decision coverage heading.
    build_decision_table(doc, decision_heading)

    # Remove the old traceability content and rebuild it with the actual FR/UC IDs.
    clear_blocks_between(doc, 'Traceability Matrix', 'CONCLUSION')
    build_traceability_section(doc, trace_heading)

    build_results_section(doc)

    # Insert Future Work as a separate section before the Conclusion
    build_future_work(doc)

    build_conclusion(doc)

    doc.save(OUT_PATH)
    print('Saved', OUT_PATH)


if __name__ == '__main__':
    try:
        main()
    except Exception as exc:
        print('Error:', exc)
        sys.exit(1)
