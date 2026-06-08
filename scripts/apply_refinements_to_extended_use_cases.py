from docx import Document
from docx.shared import Pt
import shutil

BACKUP = 'Extended use cases_original_backup.docx'
TARGET = 'Extended use cases.docx'
OUT = TARGET

# Load refined test cases from the previously generated file if available
REFINED = 'Extended use cases_refined.docx'

# Restore backup over TARGET
shutil.copyfile(BACKUP, TARGET)
print(f'Restored {TARGET} from backup {BACKUP}')

# Open target and append refinements
doc = Document(TARGET)

# Add a page break and heading
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

p = doc.add_paragraph()
run = p.add_run()
# page break
run.add_break()

h = doc.add_heading('Refinements and Additions (editor) ', level=1)

# If refined file exists, import its key tables/paragraphs
try:
    ref = Document(REFINED)
    doc.add_paragraph('The following refinements were added based on repository code inspection and to address missing/incorrect use cases. Original document content preserved above. The added items are appended below:')
    # copy paragraphs until the '7.2' heading in ref was where test cases table exists - we'll copy full refined doc body
    for element in ref.element.body:
        doc.element.body.append(element)
    print('Appended refined content from', REFINED)
except Exception as e:
    print('Could not append refined content:', e)

# Save
doc.save(OUT)
print('Saved updated document as', OUT)
