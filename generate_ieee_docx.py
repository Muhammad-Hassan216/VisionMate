from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

WORKSPACE = Path(r"d:\develop\flutter_application_1\flutter_application_1")
MD_PATH = WORKSPACE / "visionmate_ieee_paper_draft.md"
OUT_PATH = WORKSPACE / "visionmate_ieee_paper_final_camera_ready_v3.docx"
DOWNLOADS_PATH = Path(r"C:\Users\Hassan\Downloads\visionmate_ieee_paper_final_camera_ready_v3.docx")


def set_font(paragraph, size=10, bold=False, italic=False):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic


def make_two_columns(section):
    sect_pr = section._sectPr
    cols_nodes = sect_pr.xpath("./w:cols")
    cols = cols_nodes[0] if cols_nodes else sect_pr.get_or_add_cols()
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "300")


def parse_table_row(row: str) -> list[str]:
    row = row.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    return [c.strip() for c in row.split("|")]


def is_separator_row(cells: list[str]) -> bool:
    for c in cells:
        c = c.strip()
        if not c:
            continue
        if not re.fullmatch(r":?-{2,}:?", c):
            return False
    return True


def add_text_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, subscript: bool = False, superscript: bool = False):
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)
    run.bold = bold
    run.italic = italic
    run.font.subscript = subscript
    run.font.superscript = superscript
    return run


def set_code_font(paragraph):
    for run in paragraph.runs:
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(8)


def add_math_paragraph(doc, text: str):
    cleaned = text.strip()
    if cleaned == r"B=(x,y,w,h,c)":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text_run(p, "B = (x, y, w, h, c)", italic=True)
        return p

    if cleaned == r"\mathrm{IoU}(B_i,B_j)=\frac{|B_i\cap B_j|}{|B_i\cup B_j|}":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text_run(p, "IoU", italic=True)
        add_text_run(p, "(", italic=True)
        add_text_run(p, "B", italic=True)
        add_text_run(p, "i", italic=True, subscript=True)
        add_text_run(p, ", ", italic=True)
        add_text_run(p, "B", italic=True)
        add_text_run(p, "j", italic=True, subscript=True)
        add_text_run(p, ") = ", italic=True)
        add_text_run(p, "|B", italic=True)
        add_text_run(p, "i", italic=True, subscript=True)
        add_text_run(p, " ∩ B", italic=True)
        add_text_run(p, "j", italic=True, subscript=True)
        add_text_run(p, "| / |B", italic=True)
        add_text_run(p, "i", italic=True, subscript=True)
        add_text_run(p, " ∪ B", italic=True)
        add_text_run(p, "j", italic=True, subscript=True)
        add_text_run(p, "|", italic=True)
        return p

    if cleaned == r"D\approx\frac{H\cdot f}{p}":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text_run(p, "D ≈ H·f / p", italic=True)
        return p

    if cleaned == r"s=\frac{e_q\cdot e_s}{\|e_q\|\|e_s\|}":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text_run(p, "s = (e", italic=True)
        add_text_run(p, "q", italic=True, subscript=True)
        add_text_run(p, "·e", italic=True)
        add_text_run(p, "s", italic=True, subscript=True)
        add_text_run(p, ") / (‖e", italic=True)
        add_text_run(p, "q", italic=True, subscript=True)
        add_text_run(p, "‖‖e", italic=True)
        add_text_run(p, "s", italic=True, subscript=True)
        add_text_run(p, "‖)", italic=True)
        return p

    p = doc.add_paragraph(cleaned)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_run(p, cleaned, italic=True)
    return p


def build_doc():
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(1.43)
    section.right_margin = Cm(1.43)

    in_references = False
    switched_to_two_col = False
    in_code_block = False
    code_block_lines: list[str] = []

    def flush_code_block():
        nonlocal code_block_lines
        if not code_block_lines:
            return
        for code_line in code_block_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.add_run(code_line)
            set_code_font(p)
        code_block_lines = []

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if raw.startswith("```"):
            if in_code_block:
                flush_code_block()
                in_code_block = False
            else:
                in_code_block = True
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(raw)
            i += 1
            continue

        if not line or line == "---":
            i += 1
            continue

        if line == "$$":
            if i + 2 < len(lines) and lines[i + 2].strip() == "$$":
                add_math_paragraph(doc, lines[i + 1].strip())
                i += 3
                continue
            i += 1
            continue

        if line.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1

            rows = [parse_table_row(r) for r in block]
            rows = [r for r in rows if not is_separator_row(r)]
            if len(rows) >= 2:
                max_cols = max(len(r) for r in rows)
                for r in rows:
                    if len(r) < max_cols:
                        r.extend([""] * (max_cols - len(r)))

                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx, text in enumerate(row):
                        cell_p = table.cell(r_idx, c_idx).paragraphs[0]
                        cell_p.text = text
                        set_font(cell_p, size=8, bold=(r_idx == 0))
                continue

            continue

        if line.startswith("# "):
            p = doc.add_paragraph(line[2:].strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p, size=22)
            i += 1
            continue

        if line.startswith("## "):
            heading = line[3:].strip()

            if heading == "I. INTRODUCTION" and not switched_to_two_col:
                section = doc.add_section(WD_SECTION_START.CONTINUOUS)
                section.top_margin = Cm(1.9)
                section.bottom_margin = Cm(1.9)
                section.left_margin = Cm(1.43)
                section.right_margin = Cm(1.43)
                make_two_columns(section)
                switched_to_two_col = True

            if heading == "Abstract":
                p = doc.add_paragraph("Abstract-")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_font(p, size=9, bold=True)
            elif heading == "Keywords":
                p = doc.add_paragraph("Keywords-")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_font(p, size=9, bold=True)
            else:
                p = doc.add_paragraph(heading)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_font(p, size=10)

            in_references = heading == "REFERENCES"
            i += 1
            continue

        if line.startswith("### "):
            p = doc.add_paragraph(line[4:].strip())
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_font(p, size=10, italic=True)
            in_references = False
            i += 1
            continue

        if line.startswith("#### "):
            p = doc.add_paragraph(line[5:].strip())
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_font(p, size=10, italic=True)
            in_references = False
            i += 1
            continue

        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph(line.strip("*"))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p, size=11)
            i += 1
            continue

        if line.startswith("*Corresponding author:"):
            p = doc.add_paragraph(line.lstrip("*"))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p, size=9)
            i += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_font(p, size=10)
            in_references = False
            i += 1
            continue

        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if re.match(r"^\[[0-9]+\]", line) or in_references:
            pf = p.paragraph_format
            pf.left_indent = Cm(0.63)
            pf.first_line_indent = Cm(-0.63)
            set_font(p, size=8)
        else:
            set_font(p, size=10)

        i += 1

    doc.save(OUT_PATH)
    shutil.copy2(OUT_PATH, DOWNLOADS_PATH)

    ref_count = sum(1 for l in lines if re.match(r"^\[[0-9]+\]", l.strip()))
    print(f"FINAL_DOCX_CREATED: {OUT_PATH}")
    print(f"COPIED_TO_DOWNLOADS: {DOWNLOADS_PATH}")
    print(f"REFERENCE_LINES_IN_MD: {ref_count}")


if __name__ == "__main__":
    build_doc()
